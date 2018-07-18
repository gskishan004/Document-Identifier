import os
import glob
import pandas as pd
import argparse
import sys
import platform

doc_spans_multiple_page = ["Certificate of incorporation"]

CWD_PATH = os.getcwd()

p1 = os.path.join(CWD_PATH,'ml','models')
p2 = os.path.join(CWD_PATH,'ml','models','research')
p3 = os.path.join(CWD_PATH,'ml','models','research','slim')
p4 = os.path.join(CWD_PATH,'ml','models','research','object_detection')

sys.path.append(p1)
sys.path.append(p2)
sys.path.append(p3)
sys.path.append(p4)

from   docx                     import Document
from   pathlib                  import Path

from   doc_iden                 import main_doc_iden
from   lib.word_doc_iden        import main_word_doc_iden 
from   lib.word_doc_iden        import remove_doc_files
from   lib.get_move_list		import main_get_move_list
from   ml.ml_predict            import predict
from   ml.ml_train              import train
from   ml.ml_predict_multiple   import predict_multiple

def sort_doc(path,ml_flag):

    for file in glob.glob(path + '/*.jpg'):

        if(ml_flag):
            print ("Using cutom ML model to process ",file)
            output_folder, confidence = predict(file)

        else:
            print ("Using GCP APIs to process ", file)
            output_folder, confidence = main_doc_iden(file) 
            
        move_file(file, output_folder)    

def move_file(input_path, output_path):
    output_path    = os.path.join(os.getcwd(), ("output\\"+output_path))
    file_name = input_path.rsplit('\\', 1)[1]

    if not os.path.exists(output_path):
        os.makedirs(output_path)

    print ("Saving  {} to folder : {}.".format(input_path, output_path))

    os.rename(input_path, output_path+"\\"+file_name)

def move_file_to_docx(input_path, output_path):
    output_path    = os.path.join(os.getcwd(), ("output\\"+output_path+".docx"))
    file_name = input_path.rsplit('\\', 1)[1]

    myFile = Path (output_path)

    print ("Writing {} to file   : {}.".format(input_path, output_path))

    if myFile.is_file():
        document = Document(output_path)
        p = document.add_paragraph()
        r = p.add_run()
        r.add_picture(input_path)
        document.save(output_path)

    else:
        document = Document()
        p = document.add_paragraph()
        r = p.add_run()
        r.add_picture(input_path)
        document.save(output_path)

    os.remove(input_path)

if (platform.system() == "Windows"):
    os.system('cls')
else:
    os.system('clear')


parser = argparse.ArgumentParser(description='Sort the identification docs into output folder')
parser.add_argument('-i', '--input' , default="input",
					dest='i', 
					help='define the input dir')
parser.add_argument('-m', '--useML'	, action='store_true', default=False,
                    dest='m_arg',
                    help='Use ML for docuemnt prediction')
parser.add_argument('-t', '--train' , action='store_true', default=False,
                    dest='t_arg',
                    help='Train the model')
parser.add_argument('-s', '--segregate' , action='store_true', default=False,
                    dest='s_arg',
                    help='Segregate multiple docs from a image')

parser.add_argument('-w', '--word_doc' , action='store_true', default=False,
                    dest='w_arg',
                    help='Sort images from word document')

args = parser.parse_args()

if args.i == 'input':
    parser.print_help()
    print ("\nUsing default dir for input")

folder          = args.i
ml_flag         = args.m_arg
train_flag      = args.t_arg
seg_flag        = args.s_arg
word_doc_flag   = args.w_arg

image_dir       = os.path.join(os.getcwd(), (folder))

if(train_flag):
    print ("Training the model")
    train()

elif(seg_flag):
    print ("Segregating the doc")
    predict_multiple(image_dir+"\\test.jpg")

elif(word_doc_flag):
    print ("Processing images from word doc")
    #returns file names
    file_names = main_word_doc_iden(image_dir)
    file_names = sorted(file_names)

    remove_doc_files(image_dir)

    output_folder_list = []
    confidence_list    = []

    for file in file_names:
        print ("Using GCP APIs to process ", file)
        output_folder, confidence = main_doc_iden(file) 
                    
        output_folder_list.append((output_folder,file))
        confidence_list.append(confidence)

        # if there is no "No_Match" after a doc with matches category of span_multiple pages
        # save Images after No_match and untill some other doc is recognised
    final_move_list, final_move_list_doc = main_get_move_list(output_folder_list, doc_spans_multiple_page)


    for data in final_move_list:
        move_file(data[0], data[1]) 

    for data in final_move_list_doc:
        move_file_to_docx(data[0], data[1])

else:
    sort_doc(image_dir,ml_flag)